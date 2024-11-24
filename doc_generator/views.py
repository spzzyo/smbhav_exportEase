import os
from django.shortcuts import render
from django.conf import settings
from docx import Document
from docx2pdf import convert
from django.http import JsonResponse

# Ensure the static directory for PDFs exists
PDF_DIR = os.path.join(settings.BASE_DIR, "static/pdfs")
os.makedirs(PDF_DIR, exist_ok=True)

def fill_template(template_path, output_docx_path, data):
    """Fills the DOCX template with user data."""
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Save the output document
    doc.save(output_docx_path)

def index(request):
    # Extract parameters from the URL
    name = request.GET.get("name")
    importer = request.GET.get("importer")
    pod = request.GET.get("pod")  # Place of Delivery
    pol = request.GET.get("pol")  # Port of Loading
    date = request.GET.get("date")

  
    if name and importer and pod and pol and date:
     
        template_path = os.path.join(settings.BASE_DIR, "templates", "template.docx")

     
        user_data = {
            "{name}": name,
            "{importer_name}": importer,
            "{pod}": pod,
            "{pol}": pol,
            "{date}": date,
        }

        output_docx_path = os.path.join(PDF_DIR, f"{name}_document.docx")
        output_pdf_path = os.path.join(PDF_DIR, f"{name}_document.pdf")

        fill_template(template_path, output_docx_path, user_data)
        convert(output_docx_path, output_pdf_path)

        return JsonResponse({"success": True, "pdf_url": f"/static/pdfs/{name}_document.pdf"})

    return render(request, "generator/index.html", {"error": "Missing required parameters.", "pdf_path": None})



def doc_page(request):
    return render(request, "docBuddy/doc.html")

def logistics(request):
    return render(request, "logistics/logistic.html")

